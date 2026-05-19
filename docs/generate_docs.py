"""
Generador de documentación Word para el Proyecto Intermodular DAM - GranaTour
Genera el archivo Documentacion_GranaTour.docx usando python-docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from datetime import date


# ---------------------------------------------------------------------------
# Helpers de estilo
# ---------------------------------------------------------------------------

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def justify_paragraph(para):
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def center_paragraph(para):
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_justified_paragraph(doc, text, size=12, bold=False, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading1(doc, text):
    h = doc.add_heading(text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x8B, 0x1E, 0x3F)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    return h


def add_heading2(doc, text):
    h = doc.add_heading(text, level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    return h


def add_heading3(doc, text):
    h = doc.add_heading(text, level=3)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(3)
    return h


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=11)
    p.paragraph_format.space_after = Pt(3)
    return p


def style_table_header_row(table, bg_hex="8B1E3F"):
    """Aplica fondo granate y texto blanco a la primera fila de una tabla."""
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), bg_hex)
        tcPr.append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table_row(table, values, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = val
        for para in cell.paragraphs:
            para.alignment = align
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
                run.font.bold = bold
    return row


def add_page_number(doc):
    """Añade numeración de páginas en el pie de página centrado."""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.clear()

    run = footer_para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)

    run2 = footer_para.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    run2._r.append(instrText)

    run3 = footer_para.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run3._r.append(fldChar2)

    run4 = footer_para.add_run("1")
    run4.font.name = "Times New Roman"
    run4.font.size = Pt(10)

    run5 = footer_para.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run5._r.append(fldChar3)


def set_page_margins(doc, top=2.5, bottom=2.5, left=3.0, right=2.5):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


# ---------------------------------------------------------------------------
# PORTADA
# ---------------------------------------------------------------------------

def build_portada(doc):
    doc.add_paragraph()
    doc.add_paragraph()

    p_logo = doc.add_paragraph()
    center_paragraph(p_logo)
    r_logo = p_logo.add_run("🏔 GranaTour")
    r_logo.font.name = "Times New Roman"
    r_logo.font.size = Pt(36)
    r_logo.font.bold = True
    r_logo.font.color.rgb = RGBColor(0x8B, 0x1E, 0x3F)
    p_logo.paragraph_format.space_after = Pt(6)

    p_sub = doc.add_paragraph()
    center_paragraph(p_sub)
    r_sub = p_sub.add_run("Proyecto Intermodular DAM")
    r_sub.font.name = "Times New Roman"
    r_sub.font.size = Pt(22)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
    p_sub.paragraph_format.space_after = Pt(4)

    p_tipo = doc.add_paragraph()
    center_paragraph(p_tipo)
    r_tipo = p_tipo.add_run("Aplicación de Turismo y Senderismo en Granada")
    r_tipo.font.name = "Times New Roman"
    r_tipo.font.size = Pt(16)
    r_tipo.font.italic = True
    r_tipo.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    p_tipo.paragraph_format.space_after = Pt(40)

    doc.add_paragraph()
    doc.add_paragraph()

    for label, value in [
        ("Autor:", "Pablo López Huertas Moya"),
        ("Ciclo:", "Desarrollo de Aplicaciones Multiplataforma (DAM)"),
        ("Centro:", "IES Zaidín-Vergeles, Granada"),
        ("Tutor:", "Departamento de Informática"),
        ("Curso académico:", "2025-2026"),
        ("Fecha:", "Mayo 2026"),
    ]:
        p = doc.add_paragraph()
        center_paragraph(p)
        r_label = p.add_run(f"{label}  ")
        r_label.font.name = "Times New Roman"
        r_label.font.size = Pt(13)
        r_label.font.bold = True
        r_label.font.color.rgb = RGBColor(0x8B, 0x1E, 0x3F)
        r_val = p.add_run(value)
        r_val.font.name = "Times New Roman"
        r_val.font.size = Pt(13)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# ÍNDICE
# ---------------------------------------------------------------------------

def build_indice(doc):
    add_heading1(doc, "ÍNDICE DE CONTENIDOS")

    entries = [
        ("1. ANÁLISIS DEL PROBLEMA", ""),
        ("    1.1 Introducción", ""),
        ("    1.2 Objetivos del proyecto", ""),
        ("    1.3 Funciones y rendimientos deseados", ""),
        ("    1.4 Planteamiento y evaluación de soluciones alternativas", ""),
        ("    1.5 Justificación de la solución elegida", ""),
        ("    1.6 Modelado de la solución", ""),
        ("        1.6.1 Recursos humanos", ""),
        ("        1.6.2 Recursos hardware", ""),
        ("        1.6.3 Recursos software", ""),
        ("    1.7 Planificación temporal", ""),
        ("2. DISEÑO E IMPLEMENTACIÓN", ""),
        ("    2.1 Plan de empresa", ""),
        ("    2.2 Arquitectura del sistema", ""),
        ("    2.3 Tipos de usuarios y operaciones por rol", ""),
        ("    2.4 Mapa de navegación — Aplicación de escritorio", ""),
        ("    2.5 Mapa de navegación — Aplicación móvil", ""),
        ("    2.6 Diagrama Entidad-Relación", ""),
        ("    2.7 Paso a tablas", ""),
        ("    2.8 Descripción de la implementación", ""),
        ("3. FASE DE PRUEBAS", ""),
        ("    3.1 Pruebas unitarias — Aplicación de escritorio (JUnit 5)", ""),
        ("    3.2 Pruebas unitarias — Aplicación móvil (Jest)", ""),
        ("    3.3 Pruebas funcionales", ""),
        ("    3.4 Pruebas de integración", ""),
        ("4. DOCUMENTACIÓN DE LA APLICACIÓN", ""),
        ("    4.1 Introducción a las aplicaciones", ""),
        ("    4.2 Manual de instalación", ""),
        ("    4.3 Manual de usuario", ""),
        ("    4.4 Manual de administración", ""),
        ("5. CONCLUSIONES FINALES", ""),
        ("6. BIBLIOGRAFÍA", ""),
    ]

    for title, page in entries:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# CAPÍTULO 1: ANÁLISIS DEL PROBLEMA
# ---------------------------------------------------------------------------

def build_capitulo1(doc):
    add_heading1(doc, "1. ANÁLISIS DEL PROBLEMA")

    # 1.1 Introducción
    add_heading2(doc, "1.1 Introducción")
    add_justified_paragraph(doc,
        "GranaTour es un sistema informático compuesto por dos aplicaciones complementarias que "
        "tienen como objetivo centralizar la gestión y el disfrute del turismo activo y el "
        "senderismo en la provincia de Granada. La primera aplicación es una herramienta de "
        "escritorio destinada al personal administrativo y a los guías, desarrollada con JavaFX "
        "y Java 21. La segunda es una aplicación móvil multiplataforma orientada al cliente y al "
        "guía, construida con React Native y Expo sobre la plataforma Android e iOS. Ambas "
        "comparten la misma base de datos PostgreSQL alojada en Supabase, lo que garantiza la "
        "consistencia de los datos en tiempo real.", indent=True)

    add_justified_paragraph(doc,
        "La provincia de Granada cuenta con una extraordinaria riqueza natural: Sierra Nevada, "
        "el Parque Natural de la Sierra de Baza, la Alpujarra, el Geoparque de Granada y decenas "
        "de senderos homologados. Sin embargo, los turistas y senderistas locales se encuentran "
        "con información fragmentada, repartida entre webs municipales desactualizadas, "
        "aplicaciones genéricas de senderismo sin datos específicos de la zona y agencias de "
        "turismo que trabajan con sistemas propios incompatibles. Este problema dificulta la "
        "planificación de rutas, la reserva de excursiones guiadas y el seguimiento de la "
        "actividad física propia.", indent=True)

    add_justified_paragraph(doc,
        "GranaTour propone una solución integrada que concentra en una única plataforma toda la "
        "información sobre excursiones guiadas de Granada: descripción detallada, mapas "
        "interactivos con rutas GeoJSON, sistema de reservas online, tracking GPS durante la "
        "actividad, galería de fotos geolocalizadas, valoraciones de los usuarios y notificaciones "
        "push. La aplicación de escritorio complementa el sistema dotando al personal "
        "administrativo de herramientas de gestión completas: CRUD de usuarios, excursiones y "
        "reservas, generación de informes con JasperReports y una vista de control del negocio.", indent=True)

    # 1.2 Objetivos
    add_heading2(doc, "1.2 Objetivos del proyecto")

    add_justified_paragraph(doc,
        "El proyecto persigue los siguientes objetivos generales y específicos:", indent=True)

    bullets_obj = [
        "Desarrollar una aplicación móvil (Android/iOS) que permita a clientes explorar y reservar "
        "excursiones guiadas en Granada, realizando el seguimiento GPS de sus rutas.",
        "Desarrollar una aplicación de escritorio (Windows/Linux) que permita al personal "
        "administrativo gestionar usuarios, excursiones y reservas con informes exportables.",
        "Diseñar e implementar una base de datos relacional PostgreSQL en Supabase que sea compartida "
        "por ambas aplicaciones, asegurando integridad referencial y políticas RLS de seguridad.",
        "Implementar un sistema de autenticación seguro con Supabase Auth que diferencie tres roles: "
        "cliente, guía y administrador.",
        "Desarrollar el módulo de tracking GPS con grabación de ruta en tiempo real, cálculo de "
        "distancia mediante la fórmula de Haversine, desnivel positivo acumulado y exportación de "
        "la ruta en formato GeoJSON.",
        "Implementar un sistema de notificaciones push mediante Expo Notifications y Edge Functions "
        "de Supabase que informe a los clientes sobre sus reservas.",
        "Garantizar la funcionalidad básica de la app móvil en modo offline mediante caché "
        "persistente y cola de acciones pendientes.",
        "Generar informes en PDF desde la aplicación de escritorio usando JasperReports.",
        "Publicar el APK de la aplicación móvil mediante EAS Build de Expo.",
    ]
    for b in bullets_obj:
        add_bullet(doc, b)

    # 1.3 Funciones y rendimientos deseados
    add_heading2(doc, "1.3 Funciones y rendimientos deseados")

    add_justified_paragraph(doc,
        "A continuación se detallan las funcionalidades del sistema organizadas por rol de usuario:",
        indent=True)

    add_heading3(doc, "Cliente (app móvil)")
    bullets_cliente = [
        "Registro e inicio de sesión con email y contraseña.",
        "Exploración del catálogo de excursiones con filtros por zona, dificultad y precio.",
        "Búsqueda de excursiones por texto libre.",
        "Visualización del detalle de cada excursión: descripción, estadísticas, mapa con ruta, "
        "galería de fotos, datos del guía y valoraciones.",
        "Reserva de excursiones con selección de número de personas y cálculo automático del precio.",
        "Gestión de reservas propias: consulta, detalle y cancelación.",
        "Tracking GPS de rutas propias: inicio, pausa, reanudación y finalización con grabación de "
        "polilínea, cálculo de distancia, velocidad media y desnivel positivo.",
        "Historial de actividades GPS con estadísticas.",
        "Toma de fotos geolocalizadas durante el tracking o asociadas a excursiones.",
        "Publicación de valoraciones (1-5 estrellas) y comentarios sobre excursiones realizadas.",
        "Recepción de notificaciones push sobre el estado de las reservas.",
        "Funcionamiento en modo offline con sincronización al recuperar la conexión.",
        "Edición del perfil personal: nombre, bio, teléfono y foto de avatar.",
        "Visualización de estadísticas personales: km totales, excursiones y gráfico de actividad.",
    ]
    for b in bullets_cliente:
        add_bullet(doc, b)

    add_heading3(doc, "Guía (app móvil, tab adicional)")
    bullets_guia = [
        "Acceso a un panel específico con las excursiones que tiene asignadas.",
        "Visualización de la lista de reservas por excursión con nombre del cliente, "
        "número de personas y estado.",
        "Confirmación o cancelación de reservas de sus excursiones.",
        "Inicio del tracking GPS vinculado a una excursión concreta.",
        "Vista de participantes confirmados durante el seguimiento de la excursión activa.",
        "Envío de notificaciones push a los clientes al confirmar o cancelar su reserva.",
    ]
    for b in bullets_guia:
        add_bullet(doc, b)

    add_heading3(doc, "Administrador (app de escritorio)")
    bullets_admin = [
        "Gestión completa (CRUD) de usuarios: alta, consulta, modificación y baja.",
        "Gestión completa (CRUD) de excursiones: creación con imagen, coordenadas, ruta GeoJSON y "
        "asignación de guía.",
        "Gestión completa (CRUD) de reservas: visualización y cambio de estado.",
        "Generación de informes en PDF mediante JasperReports (listado de excursiones, reservas "
        "por estado, estadísticas de guías).",
        "Gestión de sesión segura con SessionManager y cierre automático por inactividad.",
        "Interfaz gráfica con JavaFX 21 y FXML con patrón MVC.",
    ]
    for b in bullets_admin:
        add_bullet(doc, b)

    add_justified_paragraph(doc,
        "En cuanto a los rendimientos esperados, se establece que la app móvil debe responder a "
        "las solicitudes de red en menos de 3 segundos en condiciones normales de conectividad, "
        "el tracking GPS debe actualizar la posición con una frecuencia mínima de 5 segundos, "
        "y la sincronización offline debe procesarse en segundo plano sin interrumpir la "
        "experiencia de usuario.", indent=True)

    # 1.4 Planteamiento y evaluación de soluciones
    add_heading2(doc, "1.4 Planteamiento y evaluación de soluciones alternativas")

    add_justified_paragraph(doc,
        "Durante la fase de diseño se evaluaron distintas alternativas técnicas antes de "
        "adoptar el stack definitivo:", indent=True)

    table_alt = doc.add_table(rows=1, cols=3)
    table_alt.style = "Table Grid"
    table_alt.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Alternativa", "Ventajas", "Inconvenientes"]
    for i, h in enumerate(headers):
        cell = table_alt.rows[0].cells[i]
        cell.text = h
    style_table_header_row(table_alt)

    rows_alt = [
        ("Aplicación web (React + Node.js)",
         "Acceso universal desde navegador; despliegue sencillo.",
         "Sin acceso a sensores nativos (GPS background, cámara nativa, haptics). "
         "Experiencia móvil inferior a una app nativa."),
        ("Solo aplicación móvil (sin desktop)",
         "Menor complejidad del proyecto.",
         "No cubre el requisito de herramienta de gestión para el administrador. "
         "Informes JasperReports no aplicables."),
        ("Base de datos MySQL local",
         "Sin coste de infraestructura en la nube.",
         "No accesible remotamente sin VPN o servidor propio. "
         "Sin autenticación integrada, sin Storage para imágenes."),
        ("Flutter en lugar de React Native",
         "Rendimiento nativo elevado; un solo lenguaje (Dart).",
         "Curva de aprendizaje nueva; ecosistema más reducido que React Native; "
         "menor integración con el stack JS del proyecto."),
        ("Firebase en lugar de Supabase",
         "Infraestructura madura y bien documentada.",
         "Base de datos NoSQL (Firestore) no relacional, lo que complica la integridad "
         "referencial. Sin soporte nativo JDBC para JavaFX. Menor control de SQL."),
        ("Spring Boot + REST API propia",
         "Control total sobre la API.",
         "Requiere desarrollar, desplegar y mantener un servidor adicional, "
         "incrementando la complejidad y el tiempo de desarrollo."),
    ]
    for vals in rows_alt:
        add_table_row(table_alt, vals)

    doc.add_paragraph()

    # 1.5 Justificación
    add_heading2(doc, "1.5 Justificación de la solución elegida")

    add_justified_paragraph(doc,
        "Tras evaluar las alternativas descritas en el apartado anterior, se optó por la "
        "combinación React Native + Expo SDK 54 para la app móvil, JavaFX 21 para la app "
        "de escritorio y Supabase (PostgreSQL) como base de datos compartida. Las razones "
        "que fundamentan esta elección son las siguientes:", indent=True)

    bullets_just = [
        "React Native con Expo permite desarrollar una aplicación nativa para Android e iOS "
        "con un único código base en TypeScript, maximizando la reutilización y reduciendo el "
        "tiempo de desarrollo.",
        "Expo SDK 54 incluye de forma nativa todas las librerías necesarias para el proyecto: "
        "expo-location para GPS, expo-notifications para push, expo-image-picker para cámara y "
        "galería, y EAS Build para la generación del APK.",
        "Supabase ofrece PostgreSQL completamente gestionado en la nube con autenticación JWT "
        "integrada, políticas RLS por fila, Storage para archivos, Edge Functions sin servidor y "
        "un SDK JavaScript oficial compatible con React Native. Además proporciona un driver JDBC "
        "compatible con Java para la app de escritorio.",
        "JavaFX 21 es la tecnología oficial de interfaces gráficas en Java y está alineada con "
        "el contenido del ciclo DAM. El uso de FXML con patrón MVC y la integración de "
        "JasperReports para informes en PDF cumplen los requisitos del módulo de Desarrollo de "
        "Interfaces.",
        "Zustand como gestor de estado en la app móvil ofrece una API sencilla sin boilerplate, "
        "compatible con middleware de persistencia mediante AsyncStorage para el modo offline.",
        "NativeWind (Tailwind CSS para React Native) acelera el desarrollo de la interfaz con "
        "un sistema de diseño coherente y soporte de dark mode mediante el prefijo dark:.",
    ]
    for b in bullets_just:
        add_bullet(doc, b)

    # 1.6 Modelado de la solución
    add_heading2(doc, "1.6 Modelado de la solución")

    add_heading3(doc, "1.6.1 Recursos humanos")
    add_justified_paragraph(doc,
        "El proyecto ha sido desarrollado íntegramente por un único desarrollador, estudiante "
        "del segundo curso del Ciclo Formativo de Grado Superior en Desarrollo de Aplicaciones "
        "Multiplataforma (DAM). El alumno ha asumido todos los roles del proyecto: análisis de "
        "requisitos, diseño de la base de datos, desarrollo backend (Supabase, Edge Functions, "
        "RPCs SQL), desarrollo frontend móvil (React Native) y desarrollo de la aplicación de "
        "escritorio (JavaFX). El proyecto ha sido tutorizado por el departamento de Informática "
        "del IES Zaidín-Vergeles de Granada.", indent=True)

    add_heading3(doc, "1.6.2 Recursos hardware")
    bullets_hw = [
        "Ordenador de desarrollo: equipo con procesador Intel Core i5/AMD Ryzen 5, 16 GB RAM, "
        "SSD 512 GB, ejecutando GNU/Linux (Fedora 41, kernel 6.17).",
        "Smartphone Android para pruebas reales (Android 13+): necesario para el testing de "
        "tracking GPS en background, notificaciones push y generación del APK.",
        "Infraestructura en la nube: servidor PostgreSQL gestionado por Supabase (plan gratuito "
        "Free Tier), con 500 MB de base de datos, 1 GB de Storage y Edge Functions.",
    ]
    for b in bullets_hw:
        add_bullet(doc, b)

    add_heading3(doc, "1.6.3 Recursos software")
    bullets_sw = [
        # Móvil
        "React Native 0.76 + Expo SDK 54 (framework móvil multiplataforma)",
        "TypeScript 5.x en modo strict (lenguaje de programación, app móvil)",
        "Expo Router 4 (navegación basada en sistema de archivos)",
        "NativeWind 4 / Tailwind CSS (estilado declarativo en React Native)",
        "Zustand 5 (gestión de estado global con middleware persist)",
        "Supabase JS SDK 2.x (@supabase/supabase-js)",
        "react-native-maps (mapas interactivos con Google Maps SDK)",
        "expo-location + expo-task-manager (GPS foreground y background)",
        "expo-notifications + expo-device (push notifications)",
        "expo-image-picker + expo-media-library (cámara y galería)",
        "expo-haptics (retroalimentación táctil háptica)",
        "@react-native-community/netinfo (detección de conectividad)",
        "EAS Build (generación de APK/AAB mediante Expo Application Services)",
        # Desktop
        "Java 21 LTS (lenguaje de programación, app de escritorio)",
        "JavaFX 21 + FXML (framework de interfaz gráfica)",
        "PostgreSQL JDBC Driver 42.x (conexión a Supabase desde Java)",
        "JasperReports 6.x (generación de informes en PDF)",
        "ControlsFX (componentes JavaFX adicionales)",
        "Ikonli (iconos vectoriales para JavaFX)",
        # Base de datos
        "Supabase (PostgreSQL 15, Auth, Storage, Edge Functions Deno)",
        "Deno (runtime Edge Functions)",
        # Herramientas
        "Visual Studio Code 1.9x (editor de código principal)",
        "IntelliJ IDEA / Eclipse (IDE para desarrollo Java/JavaFX)",
        "Git + GitHub (control de versiones y repositorio remoto)",
        "Node.js 20 LTS + npm (gestión de dependencias móvil)",
        "Expo Go / Android Emulator (testing durante desarrollo)",
    ]
    for b in bullets_sw:
        add_bullet(doc, b)

    # 1.7 Planificación temporal
    add_heading2(doc, "1.7 Planificación temporal")

    add_justified_paragraph(doc,
        "El desarrollo del proyecto se ha organizado en 13 fases iterativas (Fase 0 a Fase 12) "
        "distribuidas a lo largo de 12 semanas, desde el 25 de febrero hasta el 21 de mayo de "
        "2026. Cada fase finaliza con un commit limpio en el repositorio GitHub y una "
        "actualización del documento PROGRESS.md.", indent=True)

    table_plan = doc.add_table(rows=1, cols=4)
    table_plan.style = "Table Grid"
    table_plan.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Fase", "Nombre", "Descripción resumida", "Semana / Duración"]):
        table_plan.rows[0].cells[i].text = h
    style_table_header_row(table_plan)

    fases = [
        ("Fase 0", "Setup inicial",
         "Estructura de carpetas, cliente Supabase, tipos TypeScript, stores base, "
         "layouts Expo Router, paleta de colores.",
         "Semana 1 (2-3 días)"),
        ("Fase 1", "Autenticación",
         "Login, registro (8 campos, validación DNI/NIE), recuperación de contraseña, "
         "persistencia de sesión con AsyncStorage, protección de rutas.",
         "Semana 2 (4 días)"),
        ("Fase 2", "Explorar excursiones",
         "Catálogo con filtros (zona, dificultad, precio), búsqueda con debounce, "
         "tarjeta ExcursionCard, detalle con imagen hero y datos del guía.",
         "Semana 3 (5 días)"),
        ("Fase 3", "Sistema de reservas",
         "Modal de reserva con selector de personas y precio en tiempo real, "
         "pantalla Mis Reservas en 3 tabs, detalle y cancelación con RPC atómico.",
         "Semana 4 (5 días)"),
        ("Fase 4", "Mapa interactivo",
         "react-native-maps con markers, callouts, polyline GeoJSON y toggle "
         "lista/mapa en la pantalla Explorar.",
         "Semana 5 (4 días)"),
        ("Fase 5", "GPS Tracking",
         "Tracking foreground + background con expo-task-manager, cálculo Haversine, "
         "desnivel positivo, guardado en Supabase.",
         "Semanas 6-7 (7 días)"),
        ("Fase 6", "Fotos geolocalizadas",
         "expo-image-picker con compresión, upload a Supabase Storage, fotos "
         "durante tracking vinculadas a actividad, galería en perfil y excursión.",
         "Semana 8 (4 días)"),
        ("Fase 7", "Reviews y valoraciones",
         "CRUD de reviews con componente StarRating, media de valoraciones calculada, "
         "RPC SECURITY DEFINER para actualizar valoración del guía.",
         "Semana 8 (3 días)"),
        ("Fase 8", "Panel guía",
         "Tab condicional por rol, lista de excursiones del guía, gestión de reservas, "
         "inicio de tracking vinculado a excursión, políticas RLS.",
         "Semana 9 (5 días)"),
        ("Fase 9", "Notificaciones push",
         "expo-notifications, registro de token en BD, notificación local al reservar, "
         "recordatorio 24h, Edge Function Deno para push remoto.",
         "Semana 10 (4 días)"),
        ("Fase 10", "Modo offline",
         "NetInfo, OfflineBanner animado, caché persist con Zustand, cola de reservas "
         "pendientes con sincronización automática al reconectar.",
         "Semana 10 (4 días)"),
        ("Fase 11", "Perfil y estadísticas",
         "Edición de perfil, upload de avatar con cache-busting, stats (km, excursiones, "
         "actividades), gráfico de barras custom últimos 30 días.",
         "Semana 11 (4 días)"),
        ("Fase 12", "Polish final",
         "Onboarding 3 slides, SkeletonLoader con shimmer, haptic feedback, dark mode "
         "NativeWind, configuración EAS Build y generación de APK.",
         "Semana 12 (4 días)"),
    ]

    for vals in fases:
        add_table_row(table_plan, vals)

    doc.add_paragraph()
    doc.add_page_break()


# ---------------------------------------------------------------------------
# CAPÍTULO 2: DISEÑO E IMPLEMENTACIÓN
# ---------------------------------------------------------------------------

def build_capitulo2(doc):
    add_heading1(doc, "2. DISEÑO E IMPLEMENTACIÓN")

    # 2.1 Plan de empresa
    add_heading2(doc, "2.1 Plan de empresa")
    add_justified_paragraph(doc,
        "GranaTour S.L. es una empresa ficticia de turismo activo con sede en Granada, creada "
        "con el propósito académico de contextualizar el desarrollo del presente proyecto. La "
        "empresa se dedica a la organización y comercialización de excursiones guiadas de "
        "senderismo, escalada y turismo rural por los principales espacios naturales de la "
        "provincia de Granada.", indent=True)

    add_justified_paragraph(doc,
        "La misión de GranaTour S.L. es conectar a los amantes de la naturaleza y los turistas "
        "con las rutas y paisajes más espectaculares de Granada, ofreciendo una experiencia "
        "digital integrada que facilite la planificación, reserva y disfrute de actividades al "
        "aire libre de forma segura y cómoda.", indent=True)

    add_justified_paragraph(doc,
        "La visión empresarial consiste en convertirse en la plataforma de referencia para el "
        "turismo activo en Andalucía, expandiendo progresivamente el catálogo de rutas y la "
        "red de guías certificados, apoyándose en la tecnología para ofrecer experiencias "
        "personalizadas y sostenibles.", indent=True)

    add_justified_paragraph(doc,
        "Los valores corporativos de GranaTour S.L. son: sostenibilidad medioambiental, "
        "accesibilidad para todos los niveles físicos, seguridad en cada excursión y transparencia "
        "en precios y condiciones de cancelación.", indent=True)

    add_justified_paragraph(doc,
        "Desde el punto de vista tecnológico, GranaTour S.L. invierte en una plataforma digital "
        "propia compuesta por la aplicación móvil GranaTour App y la herramienta de gestión "
        "GranaTour Desktop, integradas ambas con la infraestructura cloud de Supabase.", indent=True)

    # 2.2 Arquitectura del sistema
    add_heading2(doc, "2.2 Arquitectura del sistema")

    add_justified_paragraph(doc,
        "El sistema sigue una arquitectura cliente-servidor en dos niveles. Ambas aplicaciones "
        "(móvil y escritorio) actúan como clientes que se comunican con los servicios de Supabase, "
        "que actúa como capa de backend. La comunicación de la app móvil se realiza a través del "
        "SDK de JavaScript de Supabase (REST/Realtime sobre HTTPS), mientras que la app de "
        "escritorio utiliza el driver JDBC de PostgreSQL para conectarse directamente a la base "
        "de datos.", indent=True)

    p_ascii = doc.add_paragraph()
    p_ascii.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ascii = p_ascii.add_run(
        "┌─────────────────────────────────────────────────────────────────────────┐\n"
        "│                        SUPABASE (Cloud)                                 │\n"
        "│  ┌─────────────┐  ┌───────────┐  ┌──────────────┐  ┌───────────────┐  │\n"
        "│  │ PostgreSQL  │  │  Auth JWT │  │   Storage    │  │ Edge Functions│  │\n"
        "│  │ (7 tablas)  │  │  (Roles)  │  │ (avatars,    │  │  (Deno/push)  │  │\n"
        "│  │             │  │           │  │  fotos)      │  │               │  │\n"
        "│  └─────────────┘  └───────────┘  └──────────────┘  └───────────────┘  │\n"
        "└───────────────────────────┬─────────────────────┬───────────────────────┘\n"
        "                            │                     │\n"
        "              Supabase JS SDK              PostgreSQL JDBC\n"
        "             (HTTPS / REST)                (TCP directo)\n"
        "                            │                     │\n"
        "         ┌──────────────────┴──┐           ┌──────┴────────────────────┐\n"
        "         │  App Móvil           │           │  App de Escritorio        │\n"
        "         │  React Native + Expo │           │  JavaFX 21 + Java 21      │\n"
        "         │  TypeScript Strict   │           │  JDBC + JasperReports     │\n"
        "         │  Android / iOS       │           │  Windows / Linux          │\n"
        "         └─────────────────────┘           └───────────────────────────┘"
    )
    r_ascii.font.name = "Courier New"
    r_ascii.font.size = Pt(8)
    p_ascii.paragraph_format.space_after = Pt(10)

    add_justified_paragraph(doc,
        "Las políticas de Row Level Security (RLS) de PostgreSQL garantizan que cada usuario "
        "solo puede acceder a los datos que le corresponden según su rol, incluso en el caso de "
        "que se comprometa la clave anónima del cliente. Las Edge Functions de Supabase, "
        "ejecutadas en el runtime Deno, se encargan del envío de notificaciones push a través "
        "de la API de Expo, accediendo a la base de datos con la clave de service_role.", indent=True)

    # 2.3 Tipos de usuarios y operaciones por rol
    add_heading2(doc, "2.3 Tipos de usuarios y operaciones por rol")

    table_roles = doc.add_table(rows=1, cols=4)
    table_roles.style = "Table Grid"
    table_roles.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Operación", "Cliente", "Guía", "Admin"]):
        table_roles.rows[0].cells[i].text = h
    style_table_header_row(table_roles)

    operaciones = [
        ("Registrarse e iniciar sesión", "Sí", "Sí", "Sí"),
        ("Explorar catálogo de excursiones", "Sí", "Sí", "Solo desktop"),
        ("Ver detalle de excursión", "Sí", "Sí", "Sí"),
        ("Crear reserva", "Sí", "No", "Sí (desktop)"),
        ("Cancelar reserva propia", "Sí", "No", "Sí (desktop)"),
        ("Ver mis reservas", "Sí", "No", "Sí (desktop)"),
        ("Confirmar/cancelar reserva de cliente", "No", "Sí (sus excursiones)", "Sí (desktop)"),
        ("Iniciar tracking GPS", "Sí", "Sí (vinculado a excursión)", "No"),
        ("Ver historial de actividades", "Sí", "Sí", "No"),
        ("Subir fotos geolocalizadas", "Sí", "Sí", "No"),
        ("Crear/editar/eliminar reviews", "Sí (si tiene reserva confirmada)", "No", "No"),
        ("Editar perfil propio", "Sí", "Sí", "Sí"),
        ("Gestionar usuarios (CRUD)", "No", "No", "Sí (desktop)"),
        ("Gestionar excursiones (CRUD)", "No", "No", "Sí (desktop)"),
        ("Gestionar reservas (CRUD)", "No", "No", "Sí (desktop)"),
        ("Generar informes PDF", "No", "No", "Sí (desktop)"),
        ("Recibir notificaciones push", "Sí", "No", "No"),
    ]
    for vals in operaciones:
        add_table_row(table_roles, vals)
    doc.add_paragraph()

    # 2.4 Mapa de navegación Desktop
    add_heading2(doc, "2.4 Mapa de navegación — Aplicación de escritorio")

    add_justified_paragraph(doc,
        "La aplicación de escritorio sigue un patrón MVC con FXML. La pantalla de login autentica "
        "al usuario y carga el layout principal, que presenta un menú lateral con cuatro secciones "
        "principales:", indent=True)

    p_nav_d = doc.add_paragraph()
    p_nav_d.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_nav_d = p_nav_d.add_run(
        "Login (LoginController.java)\n"
        "  └── Main Window (MainController.java)\n"
        "        ├── [Usuarios] (UsuariosController.java)\n"
        "        │     ├── Listar usuarios con filtros y búsqueda\n"
        "        │     ├── Crear nuevo usuario (rol, datos personales)\n"
        "        │     ├── Editar usuario seleccionado\n"
        "        │     └── Eliminar usuario (con confirmación)\n"
        "        ├── [Excursiones] (ExcursionesController.java)\n"
        "        │     ├── Listar excursiones con filtros\n"
        "        │     ├── Crear excursión (imagen, coordenadas, guía, precio)\n"
        "        │     ├── Editar excursión seleccionada\n"
        "        │     └── Eliminar excursión (con confirmación)\n"
        "        ├── [Reservas] (ReservasController.java)\n"
        "        │     ├── Listar todas las reservas\n"
        "        │     ├── Filtrar por estado (pendiente/confirmada/cancelada)\n"
        "        │     ├── Ver detalle de reserva\n"
        "        │     └── Cambiar estado de reserva\n"
        "        └── [Informes] (InformesController.java)\n"
        "              ├── Informe: listado de excursiones activas\n"
        "              ├── Informe: reservas por período\n"
        "              └── Informe: estadísticas de guías"
    )
    r_nav_d.font.name = "Courier New"
    r_nav_d.font.size = Pt(9)
    p_nav_d.paragraph_format.space_after = Pt(10)

    # 2.5 Mapa de navegación móvil
    add_heading2(doc, "2.5 Mapa de navegación — Aplicación móvil")

    add_justified_paragraph(doc,
        "La app móvil utiliza Expo Router con file-based routing. La estructura de archivos "
        "en el directorio app/ define directamente las rutas de la aplicación. El flujo de "
        "navegación es el siguiente:", indent=True)

    p_nav_m = doc.add_paragraph()
    p_nav_m.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_nav_m = p_nav_m.add_run(
        "app/index.tsx (redirect inteligente)\n"
        "  ├── ¿Onboarding hecho? → NO → app/onboarding.tsx (carrusel 3 slides)\n"
        "  ├── ¿Sesión activa? → NO → (auth)/login\n"
        "  │     ├── (auth)/register.tsx\n"
        "  │     └── (auth)/forgot-password.tsx\n"
        "  └── ¿Sesión activa? → SÍ → (tabs)\n"
        "        ├── Tab 1: (tabs)/index.tsx — Inicio\n"
        "        │     └── → excursion/[id].tsx (detalle)\n"
        "        │           └── → booking/[id].tsx (tras reservar)\n"
        "        ├── Tab 2: (tabs)/explore.tsx — Explorar\n"
        "        │     ├── Modo lista con filtros\n"
        "        │     ├── Modo mapa con markers\n"
        "        │     └── → excursion/[id].tsx\n"
        "        ├── Tab 3: (tabs)/activity.tsx — Actividad GPS\n"
        "        │     ├── Historial de actividades\n"
        "        │     └── Tracking en tiempo real (mapa + stats)\n"
        "        ├── Tab 4: (tabs)/bookings.tsx — Mis Reservas\n"
        "        │     └── → booking/[id].tsx\n"
        "        ├── Tab 5: (tabs)/profile.tsx — Perfil\n"
        "        │     ├── Edición de perfil (modal)\n"
        "        │     ├── Estadísticas + gráfico 30 días\n"
        "        │     └── Galería de fotos personal\n"
        "        └── Tab 6: (tabs)/guide.tsx — Panel Guía (solo rol 'guia')\n"
        "              └── → guide-excursion/[id].tsx\n"
        "                    ├── Lista de reservas\n"
        "                    ├── Confirmar / cancelar\n"
        "                    └── Iniciar tracking GPS vinculado"
    )
    r_nav_m.font.name = "Courier New"
    r_nav_m.font.size = Pt(9)
    p_nav_m.paragraph_format.space_after = Pt(10)

    # 2.6 Diagrama Entidad-Relación
    add_heading2(doc, "2.6 Diagrama Entidad-Relación")

    add_justified_paragraph(doc,
        "El modelo de datos de GranaTour se compone de siete entidades principales con las "
        "relaciones que se describen a continuación.", indent=True)

    add_heading3(doc, "Entidades principales")
    entidades = [
        "USUARIOS: almacena todos los usuarios del sistema independientemente del rol "
        "(cliente, guía o admin). Los atributos de estadísticas (total_km, total_excursiones, "
        "num_turnos) se calculan y actualizan desde la app. El campo supabase_auth_id vincula "
        "el registro con el sistema de autenticación de Supabase Auth.",
        "EXCURSIONES: contiene el catálogo de excursiones con toda la información geográfica "
        "(latitud, longitud, ruta_geojson en formato GeoJSON LineString) y la referencia al "
        "guía responsable (id_guia).",
        "RESERVAS: registra cada reserva realizada por un cliente para una excursión. "
        "Incluye el número de personas, el precio total calculado y el estado del ciclo "
        "de vida (pendiente → confirmada / cancelada).",
        "ACTIVIDADES: almacena los registros de tracking GPS. Puede estar vinculada "
        "opcionalmente a una excursión (cuando la inicia un guía) o ser una actividad "
        "libre del cliente. La ruta se almacena como GeoJSON en ruta_geojson.",
        "FOTOS: registra las fotos geolocalizadas subidas por los usuarios. Pueden estar "
        "asociadas a una actividad, a una excursión o a ambas. La URL apunta al bucket "
        "de Supabase Storage.",
        "REVIEWS: una por combinación (usuario, excursión), con puntuación entera de 1 a 5 "
        "y comentario opcional. Solo pueden crearse si el usuario tiene una reserva confirmada "
        "de la excursión.",
        "PUSH_TOKENS: almacena los tokens de Expo Push Notification de los dispositivos de "
        "cada usuario. Un usuario puede tener múltiples tokens (varios dispositivos). "
        "El campo activo permite invalidar tokens sin borrarlos.",
    ]
    for e in entidades:
        add_bullet(doc, e)

    add_heading3(doc, "Relaciones entre entidades")
    table_er = doc.add_table(rows=1, cols=4)
    table_er.style = "Table Grid"
    table_er.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Entidad A", "Relación", "Entidad B", "Cardinalidad"]):
        table_er.rows[0].cells[i].text = h
    style_table_header_row(table_er)
    relaciones = [
        ("USUARIOS", "guía de", "EXCURSIONES", "1:N (un guía, varias excursiones)"),
        ("USUARIOS", "realiza", "RESERVAS", "1:N (un usuario, varias reservas)"),
        ("EXCURSIONES", "tiene", "RESERVAS", "1:N (una excursión, varias reservas)"),
        ("USUARIOS", "registra", "ACTIVIDADES", "1:N"),
        ("EXCURSIONES", "asociada a (opcional)", "ACTIVIDADES", "0..1:N"),
        ("USUARIOS", "sube", "FOTOS", "1:N"),
        ("ACTIVIDADES", "contiene (opcional)", "FOTOS", "0..1:N"),
        ("EXCURSIONES", "ilustrada por (opcional)", "FOTOS", "0..1:N"),
        ("USUARIOS", "escribe", "REVIEWS", "1:N (máx. 1 por excursión)"),
        ("EXCURSIONES", "recibe", "REVIEWS", "1:N"),
        ("USUARIOS", "registra", "PUSH_TOKENS", "1:N (varios dispositivos)"),
    ]
    for r in relaciones:
        add_table_row(table_er, r)
    doc.add_paragraph()

    # 2.7 Paso a tablas
    add_heading2(doc, "2.7 Paso a tablas (modelo relacional)")

    add_justified_paragraph(doc,
        "A continuación se presenta la definición de cada tabla del modelo relacional, "
        "con sus campos, tipos, restricciones y clave. Los tipos son los de PostgreSQL.", indent=True)

    tablas_def = [
        {
            "nombre": "usuarios",
            "campos": [
                ("id_usuario", "SERIAL", "PK, NOT NULL, AUTO INCREMENT"),
                ("supabase_auth_id", "UUID", "UNIQUE, NULL (FK a Supabase Auth)"),
                ("nombre", "VARCHAR(50)", "NOT NULL"),
                ("ap1", "VARCHAR(50)", "NOT NULL"),
                ("ap2", "VARCHAR(50)", "NULL"),
                ("dni", "VARCHAR(9)", "UNIQUE, NOT NULL"),
                ("email", "VARCHAR(100)", "UNIQUE, NOT NULL"),
                ("telefono", "VARCHAR(15)", "NULL"),
                ("rol", "ENUM('cliente','guia','admin')", "NOT NULL, DEFAULT 'cliente'"),
                ("password", "VARCHAR(255)", "NULL (legacy desktop)"),
                ("avatar_url", "TEXT", "NULL"),
                ("bio", "TEXT", "NULL"),
                ("valoracion", "DECIMAL(3,2)", "NULL"),
                ("num_turnos", "INT", "NOT NULL, DEFAULT 0"),
                ("total_km", "DECIMAL(8,2)", "NOT NULL, DEFAULT 0"),
                ("total_excursiones", "INT", "NOT NULL, DEFAULT 0"),
                ("fecha_registro", "TIMESTAMPTZ", "NOT NULL, DEFAULT NOW()"),
                ("ultimo_acceso", "TIMESTAMPTZ", "NULL"),
            ],
        },
        {
            "nombre": "excursiones",
            "campos": [
                ("id_excursion", "SERIAL", "PK, NOT NULL"),
                ("nombre_ruta", "VARCHAR(100)", "NOT NULL"),
                ("zona", "VARCHAR(100)", "NOT NULL"),
                ("descripcion", "TEXT", "NULL"),
                ("duracion_horas", "DECIMAL(4,2)", "NOT NULL, CHECK > 0"),
                ("distancia_km", "DECIMAL(6,2)", "NULL"),
                ("desnivel_positivo", "INT", "NULL (metros)"),
                ("dificultad", "ENUM('facil','moderada','dificil','muy_dificil')", "NULL, DEFAULT 'moderada'"),
                ("precio_persona", "DECIMAL(6,2)", "NOT NULL, CHECK >= 0"),
                ("fecha_inicio", "DATE", "NOT NULL"),
                ("plazas_disponibles", "INT", "NOT NULL, CHECK >= 0"),
                ("id_guia", "INT", "FK → usuarios(id_usuario), NULL, ON DELETE SET NULL"),
                ("latitud", "DECIMAL(10,8)", "NULL"),
                ("longitud", "DECIMAL(11,8)", "NULL"),
                ("ruta_geojson", "TEXT", "NULL (GeoJSON LineString)"),
                ("imagen_url", "TEXT", "NULL"),
                ("activa", "BOOLEAN", "NOT NULL, DEFAULT true"),
            ],
        },
        {
            "nombre": "reservas",
            "campos": [
                ("id_reserva", "SERIAL", "PK, NOT NULL"),
                ("id_usuario", "INT", "FK → usuarios, NOT NULL, ON DELETE CASCADE"),
                ("id_excursion", "INT", "FK → excursiones, NOT NULL, ON DELETE CASCADE"),
                ("fecha_reserva", "TIMESTAMPTZ", "NOT NULL, DEFAULT NOW()"),
                ("num_personas", "INT", "NOT NULL, CHECK > 0, DEFAULT 1"),
                ("estado", "ENUM('pendiente','confirmada','cancelada')", "NOT NULL, DEFAULT 'pendiente'"),
                ("precio_total", "DECIMAL(8,2)", "NOT NULL, CHECK >= 0"),
                ("notas", "TEXT", "NULL"),
            ],
        },
        {
            "nombre": "actividades",
            "campos": [
                ("id_actividad", "SERIAL", "PK, NOT NULL"),
                ("id_usuario", "INT", "FK → usuarios, NOT NULL"),
                ("id_excursion", "INT", "FK → excursiones, NULL (tracking libre)"),
                ("fecha_inicio", "TIMESTAMPTZ", "NOT NULL"),
                ("fecha_fin", "TIMESTAMPTZ", "NULL"),
                ("distancia_km", "DECIMAL(6,2)", "NULL"),
                ("duracion_minutos", "INT", "NULL"),
                ("desnivel_positivo", "INT", "NULL (metros acumulados)"),
                ("velocidad_media", "DECIMAL(4,2)", "NULL (km/h)"),
                ("ruta_geojson", "TEXT", "NULL (array puntos GPS)"),
                ("titulo", "TEXT", "NULL"),
                ("estado", "ENUM('en_curso','completada','descartada')", "NOT NULL, DEFAULT 'en_curso'"),
            ],
        },
        {
            "nombre": "fotos",
            "campos": [
                ("id_foto", "SERIAL", "PK, NOT NULL"),
                ("id_usuario", "INT", "FK → usuarios, NOT NULL"),
                ("id_actividad", "INT", "FK → actividades, NULL"),
                ("id_excursion", "INT", "FK → excursiones, NULL"),
                ("url_storage", "TEXT", "NOT NULL"),
                ("latitud", "DECIMAL(10,8)", "NULL"),
                ("longitud", "DECIMAL(11,8)", "NULL"),
                ("descripcion", "TEXT", "NULL"),
                ("fecha", "TIMESTAMPTZ", "NOT NULL, DEFAULT NOW()"),
            ],
        },
        {
            "nombre": "reviews",
            "campos": [
                ("id_review", "SERIAL", "PK, NOT NULL"),
                ("id_usuario", "INT", "FK → usuarios, NOT NULL"),
                ("id_excursion", "INT", "FK → excursiones, NOT NULL"),
                ("puntuacion", "INT", "NOT NULL, CHECK BETWEEN 1 AND 5"),
                ("comentario", "TEXT", "NULL"),
                ("fecha", "TIMESTAMPTZ", "NOT NULL, DEFAULT NOW()"),
            ],
        },
        {
            "nombre": "push_tokens",
            "campos": [
                ("id", "SERIAL", "PK, NOT NULL"),
                ("id_usuario", "INT", "FK → usuarios, NOT NULL"),
                ("token", "TEXT", "UNIQUE, NOT NULL"),
                ("plataforma", "ENUM('ios','android')", "NOT NULL"),
                ("activo", "BOOLEAN", "NOT NULL, DEFAULT true"),
                ("fecha_registro", "TIMESTAMPTZ", "NOT NULL, DEFAULT NOW()"),
            ],
        },
    ]

    for tabla in tablas_def:
        add_heading3(doc, f"Tabla: {tabla['nombre']}")
        t = doc.add_table(rows=1, cols=3)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["Campo", "Tipo", "Restricciones"]):
            t.rows[0].cells[i].text = h
        style_table_header_row(t)
        for row in tabla["campos"]:
            add_table_row(t, row)
        doc.add_paragraph()

    # 2.8 Descripción de la implementación
    add_heading2(doc, "2.8 Descripción de la implementación")

    add_justified_paragraph(doc,
        "La implementación se ha dividido en dos aplicaciones independientes que comparten "
        "la misma base de datos. A continuación se describen los módulos principales de cada una.",
        indent=True)

    add_heading3(doc, "Módulos de la aplicación móvil")
    modulos_movil = [
        "lib/supabase.ts: inicializa el cliente de Supabase con las variables de entorno "
        "EXPO_PUBLIC_SUPABASE_URL y EXPO_PUBLIC_SUPABASE_ANON_KEY. Configura AsyncStorage "
        "como mecanismo de persistencia de sesión y activa autoRefreshToken.",
        "lib/types.ts: define las interfaces TypeScript de todas las tablas de la base de datos "
        "y los tipos auxiliares (ExcursionConGuia, ReservaConDetalles, ReviewConUsuario, PuntoGPS).",
        "lib/constants.ts: exporta la paleta de colores COLORS con las escalas primary (verde), "
        "secondary (granate), neutral y los colores funcionales.",
        "lib/utils.ts: funciones auxiliares reutilizables: formatDate, formatPrice, "
        "formatTimer, formatDistanceLabel, formatSpeedLabel, getDifficultyColor.",
        "stores/authStore.ts: gestión de sesión con Zustand persist. Almacena el usuario "
        "autenticado, el flag initializing y los métodos signIn, signUp, signOut, resetPassword.",
        "stores/excursionsStore.ts: fetch de excursiones con persist parcial (excursions, "
        "featuredExcursions, upcomingExcursions). Incluye filtros, búsqueda y fallback offline.",
        "stores/bookingsStore.ts: CRUD de reservas con RPC atómica crear_reserva_atomica. "
        "Implementa el flujo offline: si no hay conexión, encola la reserva en offlineStore.",
        "stores/activityStore.ts: motor del tracking GPS. Gestiona el estado del seguimiento, "
        "aplica el algoritmo de Haversine para el cálculo de distancias y acumula el desnivel "
        "positivo. Integra la tarea de background con expo-task-manager.",
        "stores/reviewsStore.ts: CRUD de reviews con validación de puntuación entera. Llama "
        "al RPC recalcular_valoracion_guia tras cada operación de escritura.",
        "stores/guideStore.ts: panel específico para guías. Obtiene las excursiones asignadas, "
        "sus reservas y gestiona el cambio de estado con verificación de rol.",
        "stores/pushStore.ts: registro del Expo Push Token en la tabla push_tokens mediante "
        "upsert. Gestiona la desactivación del token al cerrar sesión.",
        "stores/offlineStore.ts: cola persistente de PendingAction. Detecta la reconexión y "
        "procesa las acciones pendientes en orden secuencial.",
        "stores/profileStore.ts: actualización de datos del perfil, upload de avatar al bucket "
        "avatars con cache-busting, y estadísticas de actividad con COUNT de Supabase.",
        "app/_layout.tsx: layout raíz que inicializa el listener de sesión de Supabase, "
        "activa la detección de red con useNetworkState, registra el token push y programa "
        "los recordatorios de reservas.",
        "app/onboarding.tsx: carrusel de 3 slides para nuevos usuarios. El flag "
        "granatour_onboarding_done se persiste en AsyncStorage.",
        "components/ui/SkeletonLoader.tsx: componente base con animación shimmer (opacity "
        "interpolada 0.35-0.7) y variantes especializadas: ExcursionCardSkeleton, "
        "BookingCardSkeleton, ActivityCardSkeleton, ProfileHeaderSkeleton.",
        "components/MapView.tsx: componente ExcursionMapView reutilizable basado en "
        "react-native-maps con soporte de markers, callouts, polylines GeoJSON y modo "
        "no interactivo para uso embebido.",
        "components/ui/OfflineBanner.tsx: banner animado con fade-in/fade-out que informa "
        "del estado de conectividad y el número de acciones pendientes.",
        "supabase/functions/send-push-notification/index.ts: Edge Function en Deno que "
        "recibe una lista de id_usuario, consulta sus tokens activos y llama a la API de "
        "Expo Push Notification.",
    ]
    for m in modulos_movil:
        add_bullet(doc, m)

    add_heading3(doc, "Módulos de la aplicación de escritorio")
    modulos_desktop = [
        "model/: clases Java que mapean las entidades de la base de datos "
        "(Usuario, Excursion, Reserva) con getters y setters.",
        "dao/: capa de acceso a datos. Cada DAO contiene los métodos CRUD y "
        "ejecuta las sentencias SQL mediante PreparedStatement sobre la conexión JDBC.",
        "controller/: controladores JavaFX (LoginController, MainController, "
        "UsuariosController, ExcursionesController, ReservasController, InformesController) "
        "que reciben los eventos de la vista y delegan en los DAO.",
        "view/: archivos FXML que definen la interfaz de usuario con TableView, "
        "TextField, ComboBox y botones de acción.",
        "util/SessionManager.java: gestiona la sesión activa del administrador. "
        "Almacena el usuario autenticado en memoria y cierra la sesión por inactividad.",
        "util/DBConnection.java: clase singleton que gestiona la conexión JDBC al servidor "
        "PostgreSQL de Supabase, incluyendo el manejo del pool de conexiones.",
        "reports/: plantillas JasperReports (.jrxml) para los informes exportables en PDF. "
        "Incluye el informe de excursiones activas, reservas por período y estadísticas de guías.",
        "app/Main.java: punto de entrada de la aplicación JavaFX. Carga el stage principal "
        "e inicializa la carga del FXML de login.",
    ]
    for m in modulos_desktop:
        add_bullet(doc, m)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# CAPÍTULO 3: FASE DE PRUEBAS
# ---------------------------------------------------------------------------

def build_capitulo3(doc):
    add_heading1(doc, "3. FASE DE PRUEBAS")

    # 3.1 JUnit 5
    add_heading2(doc, "3.1 Pruebas unitarias — Aplicación de escritorio (JUnit 5)")

    add_justified_paragraph(doc,
        "Las pruebas unitarias de la aplicación de escritorio se han implementado con el "
        "framework JUnit 5 (Jupiter). Se han desarrollado clases de test para los DAO y "
        "los modelos, utilizando una base de datos de prueba separada para evitar interferir "
        "con los datos de producción.", indent=True)

    table_junit = doc.add_table(rows=1, cols=4)
    table_junit.style = "Table Grid"
    table_junit.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Clase de test", "Método probado", "Descripción", "Resultado esperado"]):
        table_junit.rows[0].cells[i].text = h
    style_table_header_row(table_junit)

    tests_junit = [
        ("UsuarioDAOTest", "findByEmail()",
         "Busca usuario existente por email",
         "Devuelve Optional<Usuario> con datos correctos"),
        ("UsuarioDAOTest", "findByEmail() — no existe",
         "Busca usuario con email inexistente",
         "Devuelve Optional.empty()"),
        ("UsuarioDAOTest", "save()",
         "Persiste un nuevo usuario en BD",
         "ID generado > 0, usuario recuperable"),
        ("UsuarioDAOTest", "update()",
         "Modifica el teléfono de un usuario",
         "Nuevo teléfono recuperable tras update"),
        ("UsuarioDAOTest", "delete()",
         "Elimina un usuario por ID",
         "findById() devuelve Optional.empty()"),
        ("ExcursionDAOTest", "findAll()",
         "Recupera todas las excursiones activas",
         "Lista no vacía, todas con activa=true"),
        ("ExcursionDAOTest", "save()",
         "Crea excursión con todos los campos",
         "ID asignado, campos iguales a los insertados"),
        ("ReservaDAOTest", "findByUsuario()",
         "Recupera reservas de un usuario",
         "Solo reservas con id_usuario correcto"),
        ("ReservaDAOTest", "updateEstado()",
         "Cambia estado de pendiente a confirmada",
         "Estado actualizado en BD"),
        ("SessionManagerTest", "setCurrentUser()",
         "Establece el usuario de sesión",
         "getCurrentUser() devuelve el mismo objeto"),
        ("SessionManagerTest", "logout()",
         "Cierra la sesión activa",
         "getCurrentUser() devuelve null"),
        ("DBConnectionTest", "getConnection()",
         "Obtiene conexión JDBC",
         "Conexión no nula, isValid() = true"),
    ]
    for r in tests_junit:
        add_table_row(table_junit, r)
    doc.add_paragraph()

    # 3.2 Jest
    add_heading2(doc, "3.2 Pruebas unitarias — Aplicación móvil (Jest)")

    add_justified_paragraph(doc,
        "Las pruebas unitarias de la aplicación móvil se han implementado con Jest y "
        "@testing-library/react-native. Se han probado las funciones utilitarias de lib/utils.ts "
        "y los helpers de los stores, garantizando su correcto funcionamiento de forma aislada "
        "sin necesidad de conexión a Supabase.", indent=True)

    table_jest = doc.add_table(rows=1, cols=4)
    table_jest.style = "Table Grid"
    table_jest.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Archivo", "Función probada", "Caso de prueba", "Resultado esperado"]):
        table_jest.rows[0].cells[i].text = h
    style_table_header_row(table_jest)

    tests_jest = [
        ("utils.test.ts", "formatDate()",
         "Fecha ISO '2026-05-20T10:00:00Z'",
         "'20 may 2026'"),
        ("utils.test.ts", "formatPrice()",
         "35.5",
         "'35,50 €'"),
        ("utils.test.ts", "formatTimer()",
         "3661 segundos",
         "'1:01:01'"),
        ("utils.test.ts", "formatDistanceLabel()",
         "12.456",
         "'12,46 km'"),
        ("utils.test.ts", "formatSpeedLabel()",
         "4.7",
         "'4,7 km/h'"),
        ("utils.test.ts", "getDifficultyColor()",
         "'dificil'",
         "Código hex del color naranja (#F97316)"),
        ("utils.test.ts", "haversineDistance()",
         "Dos coords Granada → Sierra Nevada",
         "Distancia razonable entre 25 y 40 km"),
        ("authStore.test.ts", "translateAuthError()",
         "'Invalid login credentials'",
         "'Correo o contraseña incorrectos'"),
        ("activityStore.test.ts", "addGPSPoint() filtro saltos",
         "Punto a 500m del anterior (GPS errático)",
         "El punto se descarta, distancia no aumenta"),
        ("activityStore.test.ts", "addGPSPoint() desnivel",
         "Altitud 800m → 850m → 820m",
         "desnivel_positivo = 50 (solo diff positiva)"),
        ("bookingsStore.test.ts", "calcularPrecioTotal()",
         "precio_persona=45, num_personas=3",
         "135.00"),
    ]
    for r in tests_jest:
        add_table_row(table_jest, r)
    doc.add_paragraph()

    # 3.3 Pruebas funcionales
    add_heading2(doc, "3.3 Pruebas funcionales")

    add_justified_paragraph(doc,
        "Las pruebas funcionales validan el comportamiento del sistema desde el punto de vista "
        "del usuario final, verificando que cada caso de uso produce el resultado esperado. "
        "Todas las pruebas se han ejecutado sobre la versión de producción en un dispositivo "
        "Android físico.", indent=True)

    table_func = doc.add_table(rows=1, cols=5)
    table_func.style = "Table Grid"
    table_func.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["ID", "Caso de prueba", "Entrada", "Resultado esperado", "Resultado obtenido"]):
        table_func.rows[0].cells[i].text = h
    style_table_header_row(table_func)

    pruebas_func = [
        ("PF-01", "Registro exitoso",
         "Nombre, apellidos, DNI válido, email nuevo, contraseña >= 6 chars",
         "Usuario creado, sesión iniciada, redirect a Home",
         "PASA"),
        ("PF-02", "Registro DNI duplicado",
         "DNI ya registrado en BD",
         "Mensaje de error 'El DNI ya está registrado'",
         "PASA"),
        ("PF-03", "Login correcto",
         "Email y contraseña válidos",
         "Sesión activa, redirect a Home",
         "PASA"),
        ("PF-04", "Login incorrecto",
         "Contraseña incorrecta",
         "Mensaje 'Correo o contraseña incorrectos'",
         "PASA"),
        ("PF-05", "Explorar excursiones",
         "Sin filtros",
         "Lista de todas las excursiones activas de Supabase",
         "PASA"),
        ("PF-06", "Filtrar por dificultad",
         "Filtro: dificultad='dificil'",
         "Solo excursiones con dificultad='dificil'",
         "PASA"),
        ("PF-07", "Búsqueda por texto",
         "Texto: 'Sierra Nevada'",
         "Excursiones cuyo nombre o zona contiene el texto",
         "PASA"),
        ("PF-08", "Toggle lista/mapa",
         "Botón mapa en Explorar",
         "Mapa con markers de excursiones filtradas",
         "PASA"),
        ("PF-09", "Ver detalle excursión",
         "Tap en tarjeta",
         "Pantalla con imagen, stats, mapa mini, guía, reviews",
         "PASA"),
        ("PF-10", "Reservar excursión",
         "2 personas, excursión con plazas",
         "Reserva creada, plazas decrementadas, notificación local",
         "PASA"),
        ("PF-11", "Reservar sin plazas",
         "Excursión con plazas_disponibles=0",
         "Botón Reservar deshabilitado o error al confirmar",
         "PASA"),
        ("PF-12", "Cancelar reserva",
         "Reserva propia en estado 'pendiente'",
         "Estado = 'cancelada', plazas restauradas",
         "PASA"),
        ("PF-13", "Iniciar tracking GPS",
         "Permisos concedidos, botón Iniciar",
         "Mapa activo, timer en marcha, stats actualizándose",
         "PASA"),
        ("PF-14", "Pausar y reanudar tracking",
         "Botón Pausar → Reanudar",
         "Timer parado y reanudado, distancia no cambia en pausa",
         "PASA"),
        ("PF-15", "Guardar actividad",
         "Detener → título 'Test ruta' → Guardar",
         "Actividad en BD con ruta_geojson y stats correctas",
         "PASA"),
        ("PF-16", "Subir foto durante tracking",
         "Botón cámara durante tracking activo",
         "Foto subida a Storage con coordenadas del último GPS point",
         "PASA"),
        ("PF-17", "Crear review",
         "Reserva confirmada existente, 4 estrellas",
         "Review en BD, media actualizada en ExcursionCard",
         "PASA"),
        ("PF-18", "Editar review existente",
         "Nueva puntuación: 5 estrellas",
         "Review actualizada, media recalculada",
         "PASA"),
        ("PF-19", "Eliminar review",
         "Botón Eliminar + confirmación",
         "Review eliminada, media recalculada",
         "PASA"),
        ("PF-20", "Panel guía: confirmar reserva",
         "Guía autenticado, reserva pendiente",
         "Estado = 'confirmada', push enviado al cliente",
         "PASA"),
        ("PF-21", "Editar perfil",
         "Nuevo bio y teléfono",
         "Cambios persistidos en BD y reflejados en pantalla",
         "PASA"),
        ("PF-22", "Cambiar avatar",
         "Foto de galería del dispositivo",
         "Avatar subido, URL con cache-busting, visible en perfil",
         "PASA"),
        ("PF-23", "Onboarding primera vez",
         "App instalada limpia, sin AsyncStorage",
         "Carrusel de 3 slides visible, botón Empezar lleva al login",
         "PASA"),
        ("PF-24", "Onboarding segunda vez",
         "Flag granatour_onboarding_done = true",
         "Redirect directo a login sin mostrar onboarding",
         "PASA"),
        ("PF-25", "Dark mode",
         "Sistema Android en modo oscuro",
         "Fondo oscuro, textos claros en pantalla Home",
         "PASA"),
    ]
    for r in pruebas_func:
        add_table_row(table_func, r)
    doc.add_paragraph()

    # 3.4 Pruebas de integración
    add_heading2(doc, "3.4 Pruebas de integración")

    add_justified_paragraph(doc,
        "Las pruebas de integración validan los flujos completos que atraviesan múltiples "
        "capas del sistema: store, Supabase y base de datos.", indent=True)

    add_heading3(doc, "Flujo 1: Registro → Login → Reserva → Cancelación")
    pasos_f1 = [
        "El usuario completa el formulario de registro. Se verifica que Supabase Auth "
        "crea el usuario y que el trigger handle_new_user inserta el registro en la tabla "
        "usuarios con los metadatos correctos.",
        "El usuario inicia sesión. Se verifica que authStore almacena el objeto usuario "
        "con todos los campos incluyendo el rol.",
        "El usuario navega a una excursión y realiza una reserva de 2 personas. Se "
        "verifica que el RPC crear_reserva_atomica crea la reserva y decrementa plazas "
        "en una única transacción.",
        "El usuario cancela la reserva. Se verifica que el estado cambia a 'cancelada' y "
        "que el RPC incrementar_plazas restaura las plazas disponibles.",
    ]
    for p in pasos_f1:
        add_bullet(doc, p)

    add_heading3(doc, "Flujo 2: Tracking GPS completo")
    pasos_f2 = [
        "Se solicitan permisos foreground y background con expo-location.",
        "Se inicia el tracking. La tarea de background se registra con expo-task-manager "
        "y empieza a emitir puntos GPS.",
        "Se verifica que addGPSPoint actualiza correctamente la distancia (Haversine) y "
        "el desnivel positivo acumulado en tiempo real.",
        "Se detiene el tracking y se guarda la actividad. Se verifica que la actividad "
        "aparece en la tabla actividades de Supabase con ruta_geojson, distancia, "
        "duración y desnivel correctos.",
        "Se verifica que las fotos tomadas durante el tracking quedan vinculadas a la "
        "actividad tras el guardado mediante el UPDATE de id_actividad en la tabla fotos.",
    ]
    for p in pasos_f2:
        add_bullet(doc, p)

    add_heading3(doc, "Flujo 3: Modo offline y sincronización")
    pasos_f3 = [
        "Se desactiva la conexión a internet en el dispositivo.",
        "Se verifica que el OfflineBanner aparece y que las pantallas muestran datos "
        "cacheados de excursiones y reservas.",
        "El usuario intenta crear una reserva offline. Se verifica que la reserva se "
        "encola en offlineStore.pendingActions.",
        "Se reactiva la conexión. Se verifica que el subscribe del layout detecta la "
        "transición offline → online y procesa la cola secuencialmente.",
        "Se verifica que la reserva aparece correctamente en la tabla reservas de "
        "Supabase y en la pantalla Mis Reservas.",
    ]
    for p in pasos_f3:
        add_bullet(doc, p)

    add_heading3(doc, "Flujo 4: Notificaciones push (dispositivo físico)")
    pasos_f4 = [
        "Al iniciar sesión en un dispositivo físico Android, se registra el token de "
        "Expo Push en la tabla push_tokens mediante upsert.",
        "Al crear una reserva, se verifica que el cliente recibe una notificación local "
        "inmediata (trigger: null) con el nombre de la excursión.",
        "El guía confirma la reserva desde su panel. Se verifica que la Edge Function "
        "send-push-notification es llamada y envía la notificación al cliente.",
        "Con una reserva confirmada a 24h del inicio de la excursión, se verifica que "
        "scheduleNotificationAsync programa el recordatorio con la fecha exacta.",
    ]
    for p in pasos_f4:
        add_bullet(doc, p)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# CAPÍTULO 4: DOCUMENTACIÓN DE LA APLICACIÓN
# ---------------------------------------------------------------------------

def build_capitulo4(doc):
    add_heading1(doc, "4. DOCUMENTACIÓN DE LA APLICACIÓN")

    # 4.1 Introducción
    add_heading2(doc, "4.1 Introducción a las aplicaciones")

    add_justified_paragraph(doc,
        "GranaTour se compone de dos aplicaciones independientes que comparten la misma base "
        "de datos en Supabase:", indent=True)

    add_justified_paragraph(doc,
        "La aplicación móvil GranaTour App está dirigida a clientes y guías. Permite explorar "
        "el catálogo de excursiones de la provincia de Granada, realizar y gestionar reservas, "
        "grabar rutas GPS con estadísticas en tiempo real, subir fotos geolocalizadas, dejar "
        "valoraciones y recibir notificaciones push. Está disponible para Android e iOS y se "
        "distribuye en formato APK mediante EAS Build.", indent=True)

    add_justified_paragraph(doc,
        "La aplicación de escritorio GranaTour Desktop está dirigida al personal administrativo. "
        "Proporciona herramientas de gestión completas para usuarios, excursiones y reservas, "
        "así como la generación de informes en PDF. Se distribuye como archivo JAR ejecutable "
        "con Java 21 o como instalador nativo para Windows.", indent=True)

    # 4.2 Manual de instalación
    add_heading2(doc, "4.2 Manual de instalación")

    add_heading3(doc, "Aplicación móvil — Android")
    pasos_movil = [
        "Descargar el archivo GranaTour.apk desde el enlace proporcionado (generado con "
        "EAS Build, perfil preview, plataforma Android).",
        "En el dispositivo Android, acceder a Ajustes → Seguridad → Fuentes desconocidas "
        "y activar la instalación desde fuentes desconocidas para el gestor de archivos.",
        "Abrir el archivo GranaTour.apk descargado desde el gestor de archivos del "
        "dispositivo.",
        "Seguir el asistente de instalación y aceptar los permisos solicitados: "
        "localización, cámara, galería y notificaciones.",
        "Una vez instalada, abrir la aplicación GranaTour desde el lanzador del dispositivo.",
        "En el primer lanzamiento se mostrará el onboarding de 3 slides. Pulsar 'Empezar' "
        "para acceder al login.",
        "Crear una cuenta nueva o iniciar sesión con credenciales existentes.",
    ]
    for i, p in enumerate(pasos_movil, 1):
        add_bullet(doc, f"{i}. {p}")

    add_heading3(doc, "Aplicación móvil — iOS (TestFlight)")
    pasos_ios = [
        "La distribución iOS se realiza mediante TestFlight de Apple. El administrador "
        "debe disponer de una cuenta de Apple Developer.",
        "Generar el IPA con EAS Build, perfil production, plataforma iOS.",
        "Subir el IPA a App Store Connect y distribuir mediante TestFlight a los "
        "probadores invitados por email.",
        "El usuario acepta la invitación de TestFlight e instala la app desde la "
        "aplicación TestFlight.",
    ]
    for i, p in enumerate(pasos_ios, 1):
        add_bullet(doc, f"{i}. {p}")

    add_heading3(doc, "Aplicación de escritorio — Windows/Linux")
    pasos_desktop = [
        "Verificar que Java 21 o superior está instalado en el sistema. En caso "
        "contrario, descargar desde https://adoptium.net/",
        "Descargar el archivo GranaTour-Desktop.jar desde el repositorio del proyecto.",
        "Crear el archivo de configuración config.properties en la misma carpeta que "
        "el JAR con las credenciales de acceso a Supabase:",
        "    db.url=jdbc:postgresql://db.XXXX.supabase.co:5432/postgres",
        "    db.user=postgres",
        "    db.password=[PASSWORD_DE_BD]",
        "Ejecutar la aplicación con el comando: java -jar GranaTour-Desktop.jar",
        "En Windows también se puede usar el instalador GranaTour-Setup.exe que "
        "incluye el runtime de Java.",
        "En el login, introducir las credenciales del administrador (rol='admin') "
        "registrado en la base de datos.",
    ]
    for i, p in enumerate(pasos_desktop, 1):
        add_bullet(doc, f"{i}. {p}")

    # 4.3 Manual de usuario
    add_heading2(doc, "4.3 Manual de usuario")

    add_heading3(doc, "Flujo del cliente — Primer uso")
    flujo_cliente = [
        ("Onboarding",
         "Al abrir la app por primera vez se muestra un carrusel de 3 slides que presenta "
         "las funcionalidades principales. Pulsar 'Omitir' o deslizar hasta el último slide "
         "y pulsar 'Empezar'."),
        ("Registro",
         "En la pantalla de login, pulsar '¿No tienes cuenta? Regístrate'. Completar el "
         "formulario con nombre, primer y segundo apellido, DNI/NIE, email, teléfono y "
         "contraseña (mínimo 6 caracteres). Pulsar 'Crear cuenta'."),
        ("Explorar excursiones",
         "En la tab 'Explorar', se muestra el catálogo completo. Usar la barra de búsqueda "
         "para buscar por nombre o zona. Pulsar el icono de filtro para filtrar por "
         "dificultad, zona o rango de precios. Pulsar el icono de mapa para ver los "
         "marcadores en el mapa interactivo."),
        ("Reservar excursión",
         "Desde el detalle de una excursión, pulsar 'Reservar'. En el panel inferior, "
         "seleccionar el número de personas con los botones + y -. El precio total se "
         "calcula automáticamente. Pulsar 'Confirmar reserva'. Se recibirá una notificación "
         "local de confirmación."),
        ("Ver mis reservas",
         "En la tab 'Reservas' se muestran tres pestañas: Próximas (con reservas futuras), "
         "Historial (excursiones ya realizadas) y Canceladas. Pulsar en cualquier reserva "
         "para ver el detalle. Desde el detalle se puede cancelar una reserva pendiente."),
        ("Tracking GPS",
         "En la tab 'Actividad', pulsar 'Iniciar actividad'. Conceder los permisos de "
         "localización cuando se soliciten. El mapa mostrará la ruta en tiempo real junto "
         "con el tiempo transcurrido, la distancia recorrida, la velocidad actual y el "
         "desnivel acumulado. Usar 'Pausar' y 'Reanudar' para gestionar interrupciones. "
         "Al pulsar 'Detener', se muestra el resumen con la opción de guardar con un título "
         "o descartar la actividad."),
        ("Fotos durante tracking",
         "Con el tracking activo, pulsar el botón de cámara flotante. Se puede tomar una "
         "foto nueva o seleccionar una de la galería. La foto se geolocaliza automáticamente "
         "con las coordenadas del último punto GPS."),
        ("Dejar una valoración",
         "En el detalle de una excursión, desplazarse hasta la sección 'Valoraciones'. "
         "Si se tiene una reserva confirmada de esa excursión, aparecerá el formulario para "
         "puntuar con estrellas (1-5) y escribir un comentario. Pulsar 'Publicar valoración'."),
        ("Editar perfil",
         "En la tab 'Perfil', pulsar el icono de edición junto al nombre. Se abre un modal "
         "donde se puede modificar el nombre, la bio y el teléfono. Para cambiar el avatar, "
         "pulsar sobre la foto de perfil y seleccionar una imagen de la galería."),
    ]

    for titulo, descripcion in flujo_cliente:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r1 = p.add_run(f"{titulo}: ")
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(11)
        r1.font.bold = True
        r2 = p.add_run(descripcion)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(5)

    add_heading3(doc, "Flujo del guía")
    flujo_guia = [
        ("Panel guía",
         "Los usuarios con rol 'guía' ven una tab adicional con el icono de brújula. "
         "En esta tab aparecen todas las excursiones que tienen asignadas, con el "
         "recuento de reservas pendientes y confirmadas."),
        ("Gestionar reservas",
         "Al pulsar en una excursión, se muestra la lista de clientes con reserva. "
         "Cada reserva tiene botones para confirmar (color verde) o cancelar (color rojo). "
         "Al cambiar el estado, el cliente recibe una notificación push automáticamente."),
        ("Iniciar excursión",
         "Desde la pantalla de reservas de una excursión, pulsar 'Iniciar excursión'. "
         "Se activa el tracking GPS vinculado a esa excursión. La pantalla muestra los "
         "participantes confirmados y las estadísticas del tracking en tiempo real."),
    ]
    for titulo, descripcion in flujo_guia:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r1 = p.add_run(f"{titulo}: ")
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(11)
        r1.font.bold = True
        r2 = p.add_run(descripcion)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(5)

    # 4.4 Manual de administración
    add_heading2(doc, "4.4 Manual de administración")

    add_justified_paragraph(doc,
        "La aplicación de escritorio está destinada exclusivamente al personal administrativo "
        "con rol 'admin'. Las operaciones disponibles son las siguientes:", indent=True)

    add_heading3(doc, "Gestión de usuarios")
    gestion_usuarios = [
        "Acceder a la sección 'Usuarios' desde el menú lateral.",
        "La tabla central muestra todos los usuarios registrados con nombre, email, rol y "
        "fecha de registro. Usar el campo de búsqueda para filtrar por nombre o email.",
        "Para crear un nuevo usuario, pulsar 'Nuevo usuario', completar el formulario y "
        "pulsar 'Guardar'. El usuario se inserta directamente en la tabla usuarios.",
        "Para editar un usuario, seleccionarlo en la tabla y pulsar 'Editar'. Modificar "
        "los campos deseados y confirmar con 'Guardar'.",
        "Para eliminar un usuario, seleccionarlo y pulsar 'Eliminar'. Se solicitará "
        "confirmación. La eliminación borra en cascada reservas y actividades del usuario.",
    ]
    for p in gestion_usuarios:
        add_bullet(doc, p)

    add_heading3(doc, "Gestión de excursiones")
    gestion_excursiones = [
        "Acceder a la sección 'Excursiones' desde el menú lateral.",
        "Para crear una nueva excursión, pulsar 'Nueva excursión' y completar todos los "
        "campos: nombre de la ruta, zona, descripción, duración, distancia, desnivel, "
        "dificultad, precio por persona, fecha de inicio, plazas y guía asignado.",
        "La imagen de la excursión se sube pulsando 'Seleccionar imagen'. Se almacena en "
        "Supabase Storage y la URL se guarda en el campo imagen_url.",
        "Para introducir la ruta GeoJSON, pegar el contenido del archivo .geojson en el "
        "campo habilitado. El formato debe ser un GeoJSON LineString válido.",
        "Usar el toggle 'Activa' para publicar o despublicar una excursión sin eliminarla.",
    ]
    for p in gestion_excursiones:
        add_bullet(doc, p)

    add_heading3(doc, "Gestión de reservas")
    gestion_reservas = [
        "Acceder a la sección 'Reservas' desde el menú lateral.",
        "Filtrar por estado (pendiente / confirmada / cancelada) usando el ComboBox superior.",
        "Seleccionar una reserva y pulsar 'Ver detalle' para ver la información completa "
        "(cliente, excursión, fecha, personas, precio).",
        "Cambiar el estado de una reserva con el ComboBox de estado en el panel de detalle "
        "y confirmar con 'Actualizar estado'.",
    ]
    for p in gestion_reservas:
        add_bullet(doc, p)

    add_heading3(doc, "Generación de informes")
    gestion_informes = [
        "Acceder a la sección 'Informes' desde el menú lateral.",
        "Seleccionar el tipo de informe deseado: excursiones activas, reservas por período "
        "o estadísticas de guías.",
        "Para el informe de reservas, seleccionar el rango de fechas con el DatePicker.",
        "Pulsar 'Generar informe'. Se abrirá un visor PDF con el informe generado por "
        "JasperReports.",
        "Pulsar 'Exportar PDF' para guardar el informe en el sistema de archivos local.",
    ]
    for p in gestion_informes:
        add_bullet(doc, p)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# CAPÍTULO 5: CONCLUSIONES
# ---------------------------------------------------------------------------

def build_capitulo5(doc):
    add_heading1(doc, "5. CONCLUSIONES FINALES")

    add_heading2(doc, "5.1 Grado de cumplimiento de los objetivos")

    add_justified_paragraph(doc,
        "El proyecto GranaTour ha concluido satisfactoriamente con la implementación completa "
        "de todas las fases planificadas (Fase 0 a Fase 12). La totalidad de los objetivos "
        "definidos en el apartado 1.2 han sido cumplidos:", indent=True)

    objetivos_cumplidos = [
        "La aplicación móvil multiplataforma ha sido desarrollada y probada en dispositivo "
        "Android real, con todas las funcionalidades implementadas: exploración de excursiones, "
        "reservas, tracking GPS, fotos geolocalizadas, reviews, notificaciones push, modo offline "
        "y perfil de usuario completo.",
        "La aplicación de escritorio con JavaFX proporciona el CRUD completo de usuarios, "
        "excursiones y reservas, junto con informes en PDF mediante JasperReports.",
        "La base de datos PostgreSQL en Supabase ha sido diseñada e implementada con 7 tablas, "
        "relaciones con integridad referencial, políticas RLS por rol y RPCs para operaciones "
        "atómicas.",
        "El sistema de autenticación diferencia correctamente los tres roles (cliente, guía, "
        "admin) y aplica las restricciones de acceso correspondientes tanto en la capa de datos "
        "(RLS) como en la interfaz de usuario.",
        "El tracking GPS funciona en modo foreground y background mediante expo-task-manager, "
        "con cálculo de distancia por Haversine y acumulación de desnivel positivo.",
        "Las notificaciones push funcionan en dispositivo Android físico, con Edge Function "
        "Deno para el envío remoto desde el servidor.",
        "El APK ha sido generado mediante EAS Build con el perfil preview, con "
        "bundleIdentifier com.granatour.app.",
    ]
    for o in objetivos_cumplidos:
        add_bullet(doc, o)

    add_heading2(doc, "5.2 Dificultades encontradas")

    add_justified_paragraph(doc,
        "Durante el desarrollo se encontraron diversas dificultades técnicas que "
        "requirieron investigación y adaptación de los planteamientos iniciales:", indent=True)

    dificultades = [
        ("Políticas RLS de Supabase y triggers",
         "La principal dificultad fue el comportamiento de Row Level Security en "
         "Supabase Cloud. A diferencia de PostgreSQL local, la cláusula SECURITY DEFINER "
         "en funciones y triggers no garantiza el bypass de RLS en Supabase Cloud. Fue "
         "necesario añadir políticas INSERT explícitas con WITH CHECK (true) para que el "
         "trigger handle_new_user pudiera insertar en la tabla usuarios sin violar las "
         "políticas de seguridad."),
        ("Conexión JDBC desde JavaFX a Supabase",
         "La conexión directa JDBC a PostgreSQL en Supabase Cloud requiere el uso del "
         "puerto 5432 con SSL habilitado. Fue necesario añadir parámetros SSL en la "
         "cadena de conexión JDBC (?sslmode=require) y gestionar el certificado del "
         "servidor para evitar errores de handshake TLS."),
        ("GPS en background en Android 13+",
         "Las versiones recientes de Android (13+) imponen restricciones adicionales "
         "sobre el tracking de localización en background. Fue necesario solicitar el "
         "permiso ACCESS_BACKGROUND_LOCATION de forma explícita y separada del permiso "
         "foreground, con la pantalla de permisos del sistema operativo. El fallback "
         "a watchPositionAsync en foreground garantizó el funcionamiento en dispositivos "
         "sin permiso background."),
        ("Inferencia de tipos en JOINs de Supabase",
         "El SDK JavaScript de Supabase infiere los campos de JOIN como arrays aunque "
         "sean relaciones many-to-one. Fue necesario definir tipos locales con uniones "
         "(Pick<...> | Pick<...>[] | null) y funciones de mapeo que normalizan el "
         "resultado comprobando Array.isArray."),
        ("Animación fade-out en React Native",
         "La animación de fade-out del OfflineBanner requirió gestionar el desmontaje "
         "del componente mediante un estado local. La aproximación inicial con solo "
         "Animated.timing causaba que el componente desapareciera bruscamente antes "
         "de que terminara la animación."),
        ("Detección de email duplicado en Supabase Auth",
         "Supabase Auth no devuelve un error explícito cuando se intenta registrar un "
         "email que ya existe (con confirmación de email activa). En su lugar, devuelve "
         "un usuario con identities: [] vacío. Fue necesario verificar "
         "user.identities.length === 0 para detectar este caso y mostrar el error apropiado."),
    ]
    for titulo, desc in dificultades:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r1 = p.add_run(f"{titulo}: ")
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(11)
        r1.font.bold = True
        r2 = p.add_run(desc)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(5)

    add_heading2(doc, "5.3 Propuestas de mejora futura")

    add_justified_paragraph(doc,
        "Aunque el proyecto cumple todos los objetivos planificados, se identifican las "
        "siguientes mejoras que podrían implementarse en versiones posteriores:", indent=True)

    mejoras = [
        "Chat en tiempo real entre guía y clientes de la excursión, aprovechando la "
        "funcionalidad Realtime de Supabase (suscripciones a canales PostgreSQL).",
        "Integración de un sistema de pago online (Stripe o PayPal) para completar el "
        "flujo de reserva con cobro automático.",
        "Descarga de mapas offline mediante react-native-maps con tiles locales, "
        "permitiendo la navegación sin conexión en zonas de montaña sin cobertura.",
        "Sistema de certificados de actividades: generación de un PDF personalizado "
        "al completar una excursión con las estadísticas de tracking.",
        "Panel de analítica para el administrador con gráficas de ingresos, "
        "excursiones más reservadas y guías mejor valorados.",
        "Gamificación: insignias y logros para los usuarios según los km recorridos "
        "o el número de excursiones completadas.",
        "Soporte para excursiones de varios días con tracking multi-jornada.",
        "Internacionalización (i18n) con soporte de inglés y francés para atraer "
        "turismo extranjero.",
        "Publicación de la aplicación en Google Play Store y Apple App Store.",
    ]
    for m in mejoras:
        add_bullet(doc, m)

    add_justified_paragraph(doc,
        "En términos formativos, este proyecto ha representado una oportunidad excepcional "
        "para integrar los conocimientos adquiridos a lo largo del ciclo DAM: desarrollo de "
        "aplicaciones móviles, bases de datos relacionales, servicios en la nube, patrones de "
        "diseño MVC, desarrollo de interfaces gráficas y gestión de proyectos con Git. La "
        "experiencia ha reforzado la comprensión de los desafíos reales del desarrollo de "
        "software multiplataforma, especialmente en lo relativo a la seguridad de los datos, "
        "el manejo de errores y la gestión del estado en aplicaciones con múltiples capas.", indent=True)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# CAPÍTULO 6: BIBLIOGRAFÍA
# ---------------------------------------------------------------------------

def build_capitulo6(doc):
    add_heading1(doc, "6. BIBLIOGRAFÍA")

    add_justified_paragraph(doc,
        "A continuación se recogen las fuentes de información, documentación oficial y "
        "recursos de referencia utilizados durante el desarrollo del proyecto GranaTour.",
        indent=True)

    add_heading2(doc, "6.1 Documentación oficial")

    refs = [
        ("Supabase Documentation",
         "Documentación oficial de Supabase. Cubre PostgreSQL gestionado, autenticación, "
         "Storage, Edge Functions y Row Level Security.",
         "https://supabase.com/docs"),
        ("React Native Documentation",
         "Guía oficial de React Native para el desarrollo de aplicaciones móviles "
         "multiplataforma con JavaScript y TypeScript.",
         "https://reactnative.dev/docs/getting-started"),
        ("Expo Documentation",
         "Documentación del SDK de Expo, incluyendo expo-location, expo-notifications, "
         "expo-image-picker, expo-haptics y EAS Build.",
         "https://docs.expo.dev"),
        ("Expo Router Documentation",
         "Guía de navegación basada en sistema de archivos para React Native con Expo.",
         "https://expo.github.io/router/docs"),
        ("NativeWind Documentation",
         "Documentación de NativeWind v4, implementación de Tailwind CSS para React Native.",
         "https://www.nativewind.dev/v4/overview"),
        ("Zustand GitHub Repository",
         "Repositorio oficial de Zustand con la documentación de la API y ejemplos "
         "de uso con middleware persist.",
         "https://github.com/pmndrs/zustand"),
        ("JavaFX Documentation",
         "Documentación oficial de JavaFX 21 de OpenJFX, incluyendo FXML, controles "
         "y CSS styling.",
         "https://openjfx.io/javadoc/21/"),
        ("JasperReports Documentation",
         "Guía de usuario y referencia de la API de JasperReports para la generación "
         "de informes en PDF.",
         "https://community.jaspersoft.com/documentation/jasperreports-library-ultimate-guide"),
        ("PostgreSQL Documentation",
         "Documentación oficial de PostgreSQL 15: SQL, funciones, triggers, enums "
         "y políticas de seguridad a nivel de fila (RLS).",
         "https://www.postgresql.org/docs/15/"),
        ("TypeScript Handbook",
         "Manual oficial de TypeScript con referencia del lenguaje en modo strict.",
         "https://www.typescriptlang.org/docs/handbook/intro.html"),
    ]

    for titulo, desc, url in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_titulo = p.add_run(f"{titulo}. ")
        r_titulo.font.name = "Times New Roman"
        r_titulo.font.size = Pt(11)
        r_titulo.font.bold = True
        r_desc = p.add_run(f"{desc} Disponible en: ")
        r_desc.font.name = "Times New Roman"
        r_desc.font.size = Pt(11)
        r_url = p.add_run(url)
        r_url.font.name = "Times New Roman"
        r_url.font.size = Pt(11)
        r_url.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1.25)

    add_heading2(doc, "6.2 Artículos y tutoriales")

    tutoriales = [
        ("Expo Location & Background Tasks",
         "Guía de implementación de tracking GPS en background con expo-task-manager "
         "y TaskManager.defineTask.",
         "https://docs.expo.dev/versions/latest/sdk/task-manager/"),
        ("Supabase Edge Functions with Deno",
         "Tutorial oficial para el despliegue de Edge Functions en Deno con acceso "
         "a PostgreSQL mediante service_role.",
         "https://supabase.com/docs/guides/functions"),
        ("react-native-maps GeoJSON",
         "Documentación de react-native-maps con MapView, Polyline y renderización "
         "de rutas GeoJSON.",
         "https://github.com/react-native-maps/react-native-maps"),
        ("EAS Build Documentation",
         "Guía oficial de Expo Application Services para generar APK y App Bundle "
         "con los perfiles de build.",
         "https://docs.expo.dev/build/introduction/"),
        ("PostgreSQL JDBC Driver",
         "Documentación del driver JDBC oficial de PostgreSQL para conexiones desde "
         "aplicaciones Java.",
         "https://jdbc.postgresql.org/documentation/"),
        ("Zustand Persist Middleware",
         "Documentación del middleware persist de Zustand con AsyncStorage para "
         "React Native.",
         "https://docs.pmnd.rs/zustand/integrations/persisting-store-data"),
        ("Haversine Formula Implementation",
         "Referencia matemática de la fórmula de Haversine para el cálculo de "
         "distancias sobre la esfera terrestre.",
         "https://www.movable-type.co.uk/scripts/latlong.html"),
    ]

    for titulo, desc, url in tutoriales:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_titulo = p.add_run(f"{titulo}. ")
        r_titulo.font.name = "Times New Roman"
        r_titulo.font.size = Pt(11)
        r_titulo.font.bold = True
        r_desc = p.add_run(f"{desc} Disponible en: ")
        r_desc.font.name = "Times New Roman"
        r_desc.font.size = Pt(11)
        r_url = p.add_run(url)
        r_url.font.name = "Times New Roman"
        r_url.font.size = Pt(11)
        r_url.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1.25)


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    print("Generando documentación GranaTour...")

    doc = Document()

    # Configurar página
    set_page_margins(doc)
    add_page_number(doc)

    # Construir secciones
    print("  → Portada...")
    build_portada(doc)

    print("  → Índice...")
    build_indice(doc)

    print("  → Capítulo 1: Análisis del problema...")
    build_capitulo1(doc)

    print("  → Capítulo 2: Diseño e implementación...")
    build_capitulo2(doc)

    print("  → Capítulo 3: Fase de pruebas...")
    build_capitulo3(doc)

    print("  → Capítulo 4: Documentación de la aplicación...")
    build_capitulo4(doc)

    print("  → Capítulo 5: Conclusiones finales...")
    build_capitulo5(doc)

    print("  → Capítulo 6: Bibliografía...")
    build_capitulo6(doc)

    output_path = "/home/polhm/GranaTour_APP/docs/Documentacion_GranaTour.docx"
    doc.save(output_path)
    print(f"\n✓ Documento generado correctamente en: {output_path}")

    # Verificar tamaño del archivo
    import os
    size = os.path.getsize(output_path)
    print(f"  Tamaño del archivo: {size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
